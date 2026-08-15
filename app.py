from __future__ import annotations

import io
import json
import os
import re
import sys
import urllib.error
import urllib.request
from pathlib import Path
from typing import Any

from fastapi import FastAPI, File, Form, Header, HTTPException, UploadFile
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse, Response
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field

ROOT = Path(__file__).resolve().parent
# The engine lives beside this file now. The previous release nested a whole copy of
# the WriteRoute distribution under core/, including a stale detection engine that
# reverted false-positive work when installed; that copy is gone.
sys.path.insert(0, str(ROOT))

from writeroute.audit import audit_text  # noqa: E402
from writeroute.route import suggest_text, repair_text, rewrite_with_callback, verify_text  # noqa: E402
from writeroute.contracts import compile_revision_contract  # noqa: E402
from writeroute.formatting import formatting_advice
from writeroute.optional import engine_status as _engine_status
from writeroute.genres import get_genre, load_genres  # noqa: E402

MAX_CHARS = 300_000
MAX_UPLOAD = 15 * 1024 * 1024
STATIC = ROOT / "static"
ASSETS = ROOT / "assets"

app = FastAPI(title="WriteRoute Studio", version="1.0.0")
app.mount("/static", StaticFiles(directory=STATIC), name="static")
if ASSETS.exists():
    app.mount("/assets", StaticFiles(directory=ASSETS), name="assets")



class TextPayload(BaseModel):
    text: str = Field(min_length=1, max_length=MAX_CHARS)
    genre: str = Field(min_length=2)


class SuperAuditPayload(BaseModel):
    text: str = Field(min_length=1, max_length=MAX_CHARS)
    section: str = Field(default="general")
    study_design: str | None = None
    target_guideline: str | None = None


class StatsReviewPayload(BaseModel):
    text: str = Field(min_length=1, max_length=MAX_CHARS)
    study_design: str = Field(default="observational_cohort")
    section: str = Field(default="general")


class PatternAuditPayload(BaseModel):
    text: str = Field(min_length=1, max_length=MAX_CHARS)
    section: str = Field(default="general")


class LucidLintPayload(BaseModel):
    text: str = Field(min_length=1, max_length=MAX_CHARS)


class SuggestPayload(TextPayload):
    max_candidates: int = Field(default=3, ge=1, le=5)


class VerifyPayload(BaseModel):
    original: str = Field(min_length=1, max_length=MAX_CHARS)
    candidate: str = Field(min_length=1, max_length=MAX_CHARS)
    genre: str = Field(min_length=2)


class RewritePayload(TextPayload):
    provider: str = "openai-compatible"
    model: str = ""
    base_url: str = ""
    candidates: int = Field(default=3, ge=1, le=5)
    temperature: float = Field(default=0.25, ge=0, le=1.2)


class ExportPayload(BaseModel):
    text: str = Field(max_length=MAX_CHARS)
    html: str = Field(default="", max_length=MAX_CHARS * 3)
    filename: str = "writeroute-document"
    format: str = "txt"



@app.get("/", response_class=HTMLResponse)
def home() -> str:
    return (STATIC / "index.html").read_text(encoding="utf-8")


# PDF Studio is mounted only when its extra is installed, so a prose-only install still
# starts. Visiting /pdf without it gets an explanation rather than a 404.
try:
    from pdfservice import app as pdf_app

    app.mount("/pdf", pdf_app)
    PDF_STUDIO = True
except Exception as _pdf_exc:  # missing extra, or a broken native library
    PDF_STUDIO = False
    _PDF_REASON = str(_pdf_exc)

    @app.get("/pdf", response_class=HTMLResponse)
    def pdf_unavailable() -> str:
        from writeroute.optional import install_command
        return (
            "<!doctype html><meta charset='utf-8'><title>PDF Studio</title>"
            "<body style=\"font-family:system-ui;max-width:40rem;margin:4rem auto\">"
            "<h1>PDF Studio is not installed</h1>"
            f"<p>{_PDF_REASON}</p>"
            f"<pre>{install_command('pdf')}</pre>"
            "<p><a href='/'>Back</a></p></body>"
        )


@app.get("/studio", response_class=HTMLResponse)
@app.get("/studio.html", response_class=HTMLResponse)
def studio() -> str:
    return (STATIC / "studio.html").read_text(encoding="utf-8")


@app.get("/api/health")
def health() -> dict[str, Any]:
    return {
        "status": "ok",
        "product": "WriteRoute Studio",
        "engine": "WriteRoute 2.0",
        "genres": list(load_genres().keys()),
        "byok": True,
        "persistence": "none",
        "superEngine": True,
        "engines": {name: bool(info["available"])
                    for name, info in _engine_status().items()},
        "pdfStudio": PDF_STUDIO,
    }


@app.post("/api/super-audit")
def api_super_audit(body: SuperAuditPayload) -> dict[str, Any]:
    try:
        from writeroute.super_engine import SuperEngine
        engine = SuperEngine()
        res = engine.audit_text(
            text=body.text,
            section=body.section,
            study_design=body.study_design,
            target_guideline=body.target_guideline,
        )
        return res.to_dict()
    except Exception as exc:
        raise HTTPException(500, f"Super-Audit failed: {exc}") from exc


@app.post("/api/stats-review")
def api_stats_review(body: StatsReviewPayload) -> dict[str, Any]:
    try:
        from stats_brain import ReviewContext, StatsBrainReviewer, AuctorBridge
        reviewer = StatsBrainReviewer()
        ctx = ReviewContext(
            manuscript_text=body.text,
            manifest={"study_design": body.study_design, "section": body.section},
        )
        report = reviewer.review(ctx)
        packet = AuctorBridge.packet(report)
        return {
            "report": report.to_dict(),
            "auctor_packet": packet,
        }
    except Exception as exc:
        raise HTTPException(500, f"STATS-BRAIN review failed: {exc}") from exc



@app.post("/api/pattern-audit")
def api_pattern_audit(body: PatternAuditPayload) -> dict[str, Any]:
    try:
        from scientific_pattern_engine import PatternEngine, default_lookup_path
        engine = PatternEngine.from_yaml(default_lookup_path())
        return engine.analyze(body.text, metadata={"section": body.section})
    except Exception as exc:
        raise HTTPException(500, f"Pattern audit failed: {exc}") from exc



@app.post("/api/lucid-lint")
def api_lucid_lint(body: LucidLintPayload) -> dict[str, Any]:
    try:
        from lucid_sci import LucidSciEvaluator
        evaluator = LucidSciEvaluator()
        return evaluator.evaluate(body.text)
    except Exception as exc:
        raise HTTPException(500, f"LUCID-SCI evaluation failed: {exc}") from exc


@app.get("/api/guidelines")
def api_guidelines() -> dict[str, Any]:
    try:
        from auctor_engine.guidelines import ReportingGuidelineRegistry
        reg = ReportingGuidelineRegistry()
        return {"guidelines": [g.id for g in getattr(reg, "guidelines", {}).values()] if hasattr(reg, "guidelines") else ["CONSORT", "PRISMA", "STROBE", "TRIPOD", "STARD"]}
    except Exception:
        return {"guidelines": ["CONSORT", "PRISMA", "STROBE", "TRIPOD", "STARD"]}


@app.post("/api/docx/prepare")
async def api_docx_prepare(
    file: UploadFile = File(...),
    author: str = Form(default="WriteRoute SuperEngine"),
    apply_safe_edits: bool = Form(default=True),
    track_changes: bool = Form(default=True),
    add_comments: bool = Form(default=True),
) -> Response:
    data = await file.read()
    if len(data) > MAX_UPLOAD:
        raise HTTPException(413, "file exceeds upload limit")
    
    import tempfile
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as in_f, \
         tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as out_f:
        in_path = in_f.name
        out_path = out_f.name
        in_f.write(data)

    try:
        from writeroute.super_engine import SuperEngine
        engine = SuperEngine()
        engine.process_docx(
            docx_path=in_path,
            output_path=out_path,
            author=author,
            apply_safe_edits=apply_safe_edits,
            track_changes=track_changes,
            add_comments=add_comments,
        )
        with open(out_path, "rb") as f:
            out_bytes = f.read()
        
        safe_name = f"reviewed_{file.filename or 'manuscript.docx'}"
        return Response(
            out_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{safe_name}"'},
        )
    except Exception as exc:
        raise HTTPException(500, f"DOCX preparation failed: {exc}") from exc
    finally:
        if os.path.exists(in_path): os.unlink(in_path)
        if os.path.exists(out_path): os.unlink(out_path)



@app.post("/api/audit")
def api_audit(body: TextPayload) -> dict[str, Any]:
    try:
        audit = audit_text(body.text, body.genre)
        return {
            "audit": audit.to_dict(),
            "formatting": formatting_advice(body.text, audit.genre),
        }
    except ValueError as exc:
        raise HTTPException(422, str(exc)) from exc


@app.post("/api/suggest")
def api_suggest(body: SuggestPayload) -> dict[str, Any]:
    try:
        result = suggest_text(body.text, body.genre, max_candidates=body.max_candidates)
        result["formatting"] = formatting_advice(body.text, result["genre"])
        return result
    except ValueError as exc:
        raise HTTPException(422, str(exc)) from exc


@app.post("/api/repair")
def api_repair(body: TextPayload) -> dict[str, Any]:
    try:
        return repair_text(body.text, body.genre)
    except ValueError as exc:
        raise HTTPException(422, str(exc)) from exc


@app.post("/api/verify")
def api_verify(body: VerifyPayload) -> dict[str, Any]:
    try:
        return verify_text(body.original, body.candidate, body.genre)
    except ValueError as exc:
        raise HTTPException(422, str(exc)) from exc


def _extract_upload(data: bytes, filename: str, content_type: str | None) -> tuple[str, str]:
    suffix = Path(filename or "upload.txt").suffix.lower()
    if len(data) > MAX_UPLOAD:
        raise HTTPException(413, "file exceeds 15 MB upload limit")
    if suffix in {".txt", ".md", ".markdown", ".rst", ".csv", ".log"}:
        for enc in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
            try:
                return data.decode(enc), "plain-text"
            except UnicodeDecodeError:
                continue
    if suffix == ".docx":
        try:
            from docx import Document
            doc = Document(io.BytesIO(data))
            parts: list[str] = []
            for p in doc.paragraphs:
                parts.append(p.text)
            for table in doc.tables:
                for row in table.rows:
                    parts.append("\t".join(cell.text for cell in row.cells))
            return "\n\n".join(p for p in parts if p.strip()), "docx"
        except Exception as exc:
            raise HTTPException(422, f"could not read DOCX: {exc}") from exc
    if suffix == ".pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(data))
            text = "\n\n".join((page.extract_text() or "").strip() for page in reader.pages)
            if not text.strip():
                raise ValueError("the PDF contains no extractable text")
            return text, "pdf"
        except Exception as exc:
            raise HTTPException(422, f"could not read PDF text layer: {exc}") from exc
    if suffix == ".rtf":
        raw = data.decode("latin-1", errors="replace")
        raw = re.sub(r"\\'[0-9a-fA-F]{2}", "", raw)
        raw = re.sub(r"\\[a-zA-Z]+-?\d* ?", "", raw)
        raw = raw.replace("{", "").replace("}", "")
        return raw, "rtf"
    raise HTTPException(415, "supported uploads: TXT, Markdown, DOCX, PDF with text layer, RTF, CSV")


@app.post("/api/upload")
async def api_upload(file: UploadFile = File(...), genre: str = Form(...)) -> dict[str, Any]:
    data = await file.read()
    text, source_format = _extract_upload(data, file.filename or "upload", file.content_type)
    if not text.strip():
        raise HTTPException(422, "document contains no readable text")
    if len(text) > MAX_CHARS:
        raise HTTPException(413, f"extracted document exceeds {MAX_CHARS:,} characters")
    audit = audit_text(text, genre)
    return {
        "filename": file.filename,
        "sourceFormat": source_format,
        "text": text,
        "audit": audit.to_dict(),
        "formatting": formatting_advice(text, audit.genre),
    }


def _http_json(url: str, headers: dict[str, str], payload: dict[str, Any], timeout: int = 75) -> dict[str, Any]:
    request = urllib.request.Request(
        url,
        data=json.dumps(payload).encode("utf-8"),
        headers={"Content-Type": "application/json", **headers},
        method="POST",
    )
    try:
        with urllib.request.urlopen(request, timeout=timeout) as response:
            return json.loads(response.read().decode("utf-8"))
    except urllib.error.HTTPError as exc:
        detail = exc.read().decode("utf-8", errors="replace")[:1500]
        raise RuntimeError(f"provider returned HTTP {exc.code}: {detail}") from exc
    except urllib.error.URLError as exc:
        raise RuntimeError(f"could not reach provider: {exc.reason}") from exc


def _make_byok_callback(provider: str, api_key: str, model: str, base_url: str, temperature: float):
    provider = provider.strip().lower()
    if not api_key.strip():
        raise HTTPException(400, "an API key is required for generative rewrite")
    if not model.strip():
        raise HTTPException(400, "choose or enter a model")

    def callback(contract: str, source: str) -> str:
        system = "You are WriteRoute's revision engine. Follow the editorial contract exactly. Return only the revised document text."
        user = f"EDITORIAL CONTRACT\n{contract}\n\nSOURCE DOCUMENT\n{source}"
        if provider == "anthropic":
            url = (base_url.strip().rstrip("/") or "https://api.anthropic.com") + "/v1/messages"
            data = _http_json(url, {
                "x-api-key": api_key,
                "anthropic-version": "2023-06-01",
            }, {
                "model": model,
                "max_tokens": 8192,
                "temperature": temperature,
                "system": system,
                "messages": [{"role": "user", "content": user}],
            })
            try:
                return "".join(block.get("text", "") for block in data["content"] if block.get("type") == "text").strip()
            except Exception as exc:
                raise RuntimeError(f"unexpected Anthropic response: {data}") from exc

        if provider == "deepseek":
            root = base_url.strip().rstrip("/") or "https://api.deepseek.com"
        elif provider == "openrouter":
            root = base_url.strip().rstrip("/") or "https://openrouter.ai/api/v1"
        elif provider == "openai":
            root = base_url.strip().rstrip("/") or "https://api.openai.com/v1"
        else:
            root = base_url.strip().rstrip("/")
            if not root:
                raise RuntimeError("a base URL is required for an OpenAI-compatible provider")
        url = root + "/chat/completions"
        headers = {"Authorization": f"Bearer {api_key}"}
        if provider == "openrouter":
            headers.update({"HTTP-Referer": "http://localhost", "X-Title": "WriteRoute Studio"})
        data = _http_json(url, headers, {
            "model": model,
            "temperature": temperature,
            "messages": [
                {"role": "system", "content": system},
                {"role": "user", "content": user},
            ],
        })
        try:
            content = data["choices"][0]["message"]["content"]
            if isinstance(content, list):
                content = "".join(x.get("text", "") if isinstance(x, dict) else str(x) for x in content)
            return str(content).strip()
        except Exception as exc:
            raise RuntimeError(f"unexpected compatible-provider response: {data}") from exc

    return callback


@app.post("/api/rewrite")
def api_rewrite(body: RewritePayload, x_writeroute_key: str | None = Header(default=None)) -> dict[str, Any]:
    callback = _make_byok_callback(body.provider, x_writeroute_key or "", body.model, body.base_url, body.temperature)
    try:
        return rewrite_with_callback(body.text, callback, body.genre, candidates=body.candidates)
    except ValueError as exc:
        raise HTTPException(422, str(exc)) from exc
    except RuntimeError as exc:
        raise HTTPException(502, str(exc)) from exc


@app.post("/api/export")
def api_export(body: ExportPayload) -> Response:
    safe = re.sub(r"[^A-Za-z0-9._-]+", "-", body.filename).strip("-.") or "writeroute-document"
    fmt = body.format.lower()
    if fmt == "txt":
        return Response(body.text, media_type="text/plain; charset=utf-8", headers={"Content-Disposition": f'attachment; filename="{safe}.txt"'})
    if fmt == "md":
        return Response(body.text, media_type="text/markdown; charset=utf-8", headers={"Content-Disposition": f'attachment; filename="{safe}.md"'})
    if fmt == "html":
        content = body.html or "<p>" + body.text.replace("&", "&amp;").replace("<", "&lt;").replace("\n", "<br>") + "</p>"
        doc = f"<!doctype html><meta charset='utf-8'><title>{safe}</title><body>{content}</body>"
        return Response(doc, media_type="text/html; charset=utf-8", headers={"Content-Disposition": f'attachment; filename="{safe}.html"'})
    if fmt == "docx":
        try:
            from docx import Document
            from docx.shared import Pt
            doc = Document()
            styles = doc.styles
            styles["Normal"].font.name = "Aptos"
            styles["Normal"].font.size = Pt(11)
            for block in re.split(r"\n\s*\n", body.text):
                if block.strip():
                    doc.add_paragraph(block.strip())
            out = io.BytesIO(); doc.save(out)
            return Response(out.getvalue(), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f'attachment; filename="{safe}.docx"'})
        except Exception as exc:
            raise HTTPException(500, f"DOCX export failed: {exc}") from exc
    if fmt in ("tex", "latex"):
        try:
            from writeroute.latex_export import markdown_to_latex
            doc = markdown_to_latex(body.text, title=safe)
            return Response(doc, media_type="text/x-tex; charset=utf-8", headers={"Content-Disposition": f'attachment; filename="{safe}.tex"'})
        except Exception as exc:
            raise HTTPException(500, f"LaTeX export failed: {exc}") from exc
    raise HTTPException(422, "format must be txt, md, html, docx, or tex")


# ------------------------------------------------------------------ Overleaf LaTeX Suite

class LatexPreviewPayload(BaseModel):
    text: str = Field(min_length=1, max_length=MAX_CHARS)
    title: str = "Manuscript"
    doc_class: str = "article"
    author: str = "Author"


@app.post("/api/latex/preview")
def api_latex_preview(body: LatexPreviewPayload) -> dict[str, str]:
    """Generates clean, syntax-highlighted LaTeX source code from Markdown/rich editor text."""
    from writeroute.latex_export import markdown_to_latex
    tex = markdown_to_latex(body.text, title=body.title, doc_class=body.doc_class, author=body.author)
    return {"latex": tex, "doc_class": body.doc_class, "title": body.title}


# ------------------------------------------------------------------ Adobe PDF Manipulation Suite
@app.post("/api/pdf/merge")
async def api_pdf_merge(files: list[UploadFile] = File(...)) -> Response:
    """Merges multiple uploaded PDF documents into a single PDF."""
    if len(files) < 2:
        raise HTTPException(400, "Provide at least two PDF files to merge.")
    from writeroute.pdf_tools import merge_pdfs
    pdf_bytes_list = []
    for f in files:
        data = await f.read()
        if not data.startswith(b"%PDF"):
            raise HTTPException(400, f"File {f.filename} is not a valid PDF.")
        pdf_bytes_list.append(data)

    merged = merge_pdfs(pdf_bytes_list)
    return Response(
        merged,
        media_type="application/pdf",
        headers={"Content-Disposition": 'attachment; filename="merged_document.pdf"'},
    )


@app.post("/api/pdf/split")
async def api_pdf_split(file: UploadFile = File(...), pages: str = Form("1")) -> Response:
    """Extracts specific page ranges (e.g. '1, 3, 5-8') from a PDF."""
    data = await file.read()
    if not data.startswith(b"%PDF"):
        raise HTTPException(400, "File is not a valid PDF.")
    from writeroute.pdf_tools import split_pdf
    # Parse page specification (1-indexed user strings -> 0-indexed ints)
    indices = []
    for part in pages.split(","):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            p_start, p_end = part.split("-", 1)
            try:
                s_int, e_int = int(p_start.strip()), int(p_end.strip())
                for p_num in range(s_int, e_int + 1):
                    indices.append(p_num - 1)
            except ValueError:
                continue
        else:
            try:
                indices.append(int(part) - 1)
            except ValueError:
                continue

    if not indices:
        raise HTTPException(400, "No valid page numbers provided.")

    extracted = split_pdf(data, indices)
    return Response(
        extracted,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="extracted_{file.filename or "pages"}.pdf"'},
    )


@app.post("/api/pdf/rotate")
async def api_pdf_rotate(
    file: UploadFile = File(...),
    angle: int = Form(90),
    pages: str | None = Form(None),
) -> Response:
    """Rotates pages of a PDF by 90, 180, or 270 degrees."""
    data = await file.read()
    if not data.startswith(b"%PDF"):
        raise HTTPException(400, "File is not a valid PDF.")
    from writeroute.pdf_tools import rotate_pdf
    page_indices = None
    if pages and pages.strip():
        page_indices = []
        for part in pages.split(","):
            part = part.strip()
            if "-" in part:
                p_start, p_end = part.split("-", 1)
                try:
                    s_int, e_int = int(p_start.strip()), int(p_end.strip())
                    for p_num in range(s_int, e_int + 1):
                        page_indices.append(p_num - 1)
                except ValueError:
                    continue
            else:
                try:
                    page_indices.append(int(part) - 1)
                except ValueError:
                    continue

    rotated = rotate_pdf(data, angle=angle, page_indices=page_indices)
    return Response(
        rotated,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="rotated_{file.filename or "doc"}.pdf"'},
    )


@app.post("/api/pdf/watermark")
async def api_pdf_watermark(
    file: UploadFile = File(...),
    text: str = Form("CONFIDENTIAL"),
    opacity: float = Form(0.25),
    angle: float = Form(45.0),
) -> Response:
    """Stamps a diagonal custom watermark on a PDF."""
    data = await file.read()
    if not data.startswith(b"%PDF"):
        raise HTTPException(400, "File is not a valid PDF.")
    from writeroute.pdf_tools import watermark_pdf
    watermarked = watermark_pdf(data, text=text, opacity=opacity, angle=angle)
    return Response(
        watermarked,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="watermarked_{file.filename or "doc"}.pdf"'},
    )


@app.post("/api/pdf/redact")
async def api_pdf_redact(
    file: UploadFile = File(...),
    terms: str = Form(...),
) -> Response:
    """Permanently redacts sensitive terms from a PDF."""
    data = await file.read()
    if not data.startswith(b"%PDF"):
        raise HTTPException(400, "File is not a valid PDF.")
    from writeroute.pdf_tools import redact_pdf
    term_list = [t.strip() for t in terms.split(",") if t.strip()]
    if not term_list:
        raise HTTPException(400, "Provide at least one search term to redact.")

    redacted = redact_pdf(data, term_list)
    return Response(
        redacted,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="redacted_{file.filename or "doc"}.pdf"'},
    )


@app.post("/api/pdf/extract")
async def api_pdf_extract(file: UploadFile = File(...)) -> dict[str, Any]:
    """Extracts semantic text, metadata, and per-page metrics from an uploaded PDF."""
    data = await file.read()
    if not data.startswith(b"%PDF"):
        raise HTTPException(400, "File is not a valid PDF.")
    from writeroute.pdf_tools import extract_pdf_semantic
    return extract_pdf_semantic(data)


# ------------------------------------------------------------------ DocXML & Auctor Doctrine Suite
@app.post("/api/docx/import-rich")
async def api_docx_import_rich(file: UploadFile = File(...)) -> dict[str, Any]:
    """Parses a Word DOCX OOXML package, extracting rich HTML formatting, headings, tables, SDTs and comments."""
    data = await file.read()
    if not data.startswith(b"PK"):
        raise HTTPException(400, "File is not a valid DOCX package.")
    from writeroute.docxml_engine import docx_to_rich_html
    return docx_to_rich_html(data)


class DoctrineAuditPayload(BaseModel):
    original_text: str = Field(min_length=1, max_length=MAX_CHARS)
    candidate_text: str | None = None
    authority: str = "substantive"  # mechanical, copyedit, substantive, developmental


@app.post("/api/auctor/doctrine-audit")
def api_auctor_doctrine_audit(body: DoctrineAuditPayload) -> dict[str, Any]:
    """Evaluates text against the Auctor Writing Doctrine: Fact Ledger invariants,

    three-channel leakage separation, and declared revision authority compliance.
    """
    from writeroute.auctor_doctrine import FactLedger, ThreeChannelEnforcer, RevisionAuthority
    ledger = FactLedger.extract_from_text(body.original_text)
    channel_leaks = ThreeChannelEnforcer.check_channel_leakage(body.original_text)
    clean_substantive = ThreeChannelEnforcer.sanitize_substantive_channel(body.original_text)

    authority_check = {"valid": True}
    if body.candidate_text:
        authority_check = RevisionAuthority.validate_action(
            body.authority,
            body.original_text,
            body.candidate_text,
        )

    return {
        "fact_ledger": {
            "numbers_count": len(ledger.numbers),
            "numbers": ledger.numbers[:15],
            "percentages": ledger.percentages,
            "confidence_intervals": ledger.confidence_intervals,
            "p_values": ledger.p_values,
            "effect_measures": ledger.effect_measures,
            "directions": ledger.directions[:10],
            "negations": ledger.negations[:10],
            "identifiers": ledger.identifiers,
            "citations": ledger.citations,
            "table_fig_refs": ledger.table_fig_refs,
        },
        "three_channel_status": {
            "is_clean": len(channel_leaks) == 0,
            "leaks": channel_leaks,
            "sanitized_preview": clean_substantive[:300] + "..." if len(clean_substantive) > 300 else clean_substantive,
        },
        "authority_validation": authority_check,
    }


@app.get("/api/router/lead-skill")
def api_router_lead_skill(artifact: str = "docx") -> dict[str, Any]:
    """Returns the lead skill routing for a given artifact type."""
    from writeroute.auctor_doctrine import document_artifact_router
    return document_artifact_router(artifact)


# ------------------------------------------------------------------ Native Citation Manager Suite
class CitationParsePayload(BaseModel):
    raw_text: str = Field(min_length=1)
    format: str = "bibtex"  # bibtex, ris


@app.post("/api/citations/parse")
def api_citations_parse(body: CitationParsePayload) -> dict[str, Any]:
    """Parses raw BibTeX or RIS strings into structured reference items."""
    from writeroute.citation_engine import parse_bibtex, parse_ris
    if body.format.lower() == "ris":
        items = parse_ris(body.raw_text)
    else:
        items = parse_bibtex(body.raw_text)
    return {"count": len(items), "items": [item.to_dict() for item in items]}


class CitationFormatPayload(BaseModel):
    items: list[dict[str, Any]] | None = None
    bibtex: str | None = None
    style: str = "apa"  # apa, vancouver, nature, ieee, chicago


@app.post("/api/citations/format")
def api_citations_format(body: CitationFormatPayload) -> dict[str, Any]:
    """Formats in-text citations and full bibliography entries from reference items or raw BibTeX."""
    from writeroute.citation_engine import ReferenceItem, CitationFormatter, parse_bibtex
    
    if body.items:
        ref_items = [ReferenceItem.from_dict(d) for d in body.items]
    elif body.bibtex:
        ref_items = parse_bibtex(body.bibtex)
    else:
        ref_items = []

    entries = []
    legacy_citations = []
    for idx, item in enumerate(ref_items, 1):
        in_t = CitationFormatter.format_in_text([item], style=body.style, indices=[idx])
        bib_e = CitationFormatter.format_bibliography_entry(item, style=body.style, index=idx)
        
        # Format legacy in_text specifically for style expectations
        if body.style == "nature":
            leg_in_t = f"^{{{idx}}}"
        elif body.style in ("vancouver", "ieee"):
            leg_in_t = f"[{idx}]"
        elif body.style == "chicago":
            leg_in_t = f"({item.authors[0].family} {item.year})" if item.authors else f"({item.cite_key})"
        else:  # APA
            leg_in_t = f"({item.authors[0].family}, {item.year})" if item.authors else f"({item.cite_key})"

        auth_joined = " and ".join(f"{a.family}, {a.given}".strip(", ") for a in item.authors) if item.authors else "Unknown Author"
        leg_bib = f"{auth_joined} ({item.year}). {item.title}. <em>{item.journal}</em>, {item.volume}, {item.pages}."

        entries.append({
            "id": item.id,
            "cite_key": item.cite_key,
            "in_text": in_t,
            "bibliography_entry": bib_e,
        })
        legacy_citations.append({
            "key": item.cite_key,
            "in_text": leg_in_t,
            "full_citation": leg_bib,
        })



    grouped_in_text = CitationFormatter.format_in_text(ref_items, style=body.style)
    return {
        "count": len(ref_items),
        "style": body.style,
        "grouped_in_text": grouped_in_text,
        "entries": entries,
        "citations": legacy_citations,
    }



class CitationExportPayload(BaseModel):
    items: list[dict[str, Any]]
    format: str = "bibtex"  # bibtex, ris


@app.post("/api/citations/export")
def api_citations_export(body: CitationExportPayload) -> dict[str, Any]:
    """Exports reference library items into BibTeX or RIS string."""
    from writeroute.citation_engine import ReferenceItem, export_library_to_bibtex, export_library_to_ris
    ref_items = [ReferenceItem.from_dict(d) for d in body.items]
    if body.format.lower() == "ris":
        data = export_library_to_ris(ref_items)
        ext = "ris"
    else:
        data = export_library_to_bibtex(ref_items)
        ext = "bib"
    return {"format": ext, "content": data, "count": len(ref_items)}






# ------------------------------------------------------------------ Writing Master (AIWD) Engine

class AiwdScanPayload(BaseModel):
    text: str = Field(min_length=1, max_length=MAX_CHARS)
    genre: str = "academic"


@app.post("/api/aiwd/scan")
def api_aiwd_scan(body: AiwdScanPayload) -> dict[str, Any]:
    """Runs Writing-Master ontology detection and anti-slop feature extraction."""
    from aiwd.skillengine import SkillRegistry
    from aiwd.scoring import scan_text
    from aiwd.textmodel import parse
    from aiwd.rewrite import suggest

    reg = SkillRegistry.load()
    rep = scan_text(body.text, registry=reg, genre=body.genre)
    sample = parse(body.text)
    suggs = suggest(sample, reg)

    return {
        "detectionResult": rep.get("detectionResult", {}),
        "tokenCount": rep.get("tokenCount", 0),
        "sentenceCount": rep.get("sentenceCount", 0),
        "paragraphCount": rep.get("paragraphCount", 0),
        "features": rep.get("features", []),
        "allowListExemptions": rep.get("allowListExemptions", []),
        "reportedVoiceDiscounts": rep.get("reportedVoiceDiscounts", []),
        "reportedVoiceFraction": rep.get("reportedVoiceFraction", 0.0),
        "suggestions": [
            {
                "start": s.start,
                "end": s.end,
                "original": s.original,
                "options": s.options,
                "feature_id": s.feature_id,
                "family": s.family,
                "rationale": s.rationale,
                "safe": s.safe,
            }
            for s in suggs
        ],
    }


class AiwdCleanPayload(BaseModel):
    text: str = Field(min_length=1, max_length=MAX_CHARS)


@app.post("/api/aiwd/clean")
def api_aiwd_clean(body: AiwdCleanPayload) -> dict[str, Any]:
    """Applies safe de-slop replacements and verifies semantic preservation gate."""
    from aiwd.skillengine import SkillRegistry
    from aiwd.textmodel import parse
    from aiwd.rewrite import suggest, apply_safe
    from aiwd.revision import preservation_gate

    reg = SkillRegistry.load()
    sample = parse(body.text)
    suggs = suggest(sample, reg)
    cleaned_text, applied_count = apply_safe(body.text, suggs)
    passes_gate, violations = preservation_gate(body.text, cleaned_text)

    return {
        "original_text": body.text,
        "cleaned_text": cleaned_text,
        "applied_count": applied_count,
        "passes_preservation_gate": passes_gate,
        "gate_violations": violations,
    }


@app.get("/api/aiwd/packs")
def api_aiwd_packs() -> dict[str, Any]:
    """Lists loaded skill packs and feature definitions."""
    from aiwd.skillengine import SkillRegistry
    reg = SkillRegistry.load()
    return {
        "packs": reg.packs,
        "families": reg.families,
        "feature_count": len(reg.features),
    }


@app.exception_handler(Exception)
async def generic_error(_request, exc: Exception):
    if isinstance(exc, HTTPException):
        raise exc
    return JSONResponse(status_code=500, content={"detail": str(exc)})


app.mount("/", StaticFiles(directory=STATIC, html=True), name="static_root")



from __future__ import annotations

import io
import json
import threading
import sys
from pathlib import Path
from http.server import BaseHTTPRequestHandler, HTTPServer

from docx import Document
from fastapi.testclient import TestClient

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))
import app as studio

client = TestClient(studio.app)
SLOP = (
    "Here is a revised version that provides a comprehensive overview. "
    "It is important to note that studies show this robust solution can transform outcomes. "
    "In conclusion, this marks a pivotal moment."
)


def make_docx(text: str) -> bytes:
    doc = Document()
    doc.add_heading("Test report", level=1)
    doc.add_paragraph(text)
    out = io.BytesIO(); doc.save(out)
    return out.getvalue()


def test_health_and_ui():
    health = client.get("/api/health")
    assert health.status_code == 200
    assert health.json()["byok"] is True
    landing = client.get("/")
    assert landing.status_code == 200
    assert "Named editorial findings" in landing.text
    # The landing page must not claim an authorship verdict, which is the one thing the
    # engine refuses to produce.
    assert "AI-generated" not in landing.text
    studio = client.get("/studio")
    assert studio.status_code == 200
    assert "contenteditable=\"true\"" in studio.text
    # Genre is a required choice, not a default guess.
    assert 'value="" selected disabled' in studio.text


def test_docx_upload_audit_suggest_repair_verify_export():
    upload = client.post(
        "/api/upload",
        data={"genre": "professional-report"},
        files={"file": ("report.docx", make_docx(SLOP), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert upload.status_code == 200, upload.text
    data = upload.json()
    assert data["sourceFormat"] == "docx"
    assert data["audit"]["counts"]["findings"] >= 3
    assert data["formatting"]["genre"] == "professional-report"

    text = data["text"]
    suggest = client.post("/api/suggest", json={"text": text, "genre": "professional-report", "max_candidates": 3})
    assert suggest.status_code == 200
    assert suggest.json()["findings"]

    repair = client.post("/api/repair", json={"text": text, "genre": "professional-report"})
    assert repair.status_code == 200
    repaired = repair.json()["finalText"]
    assert repair.json()["changed"] is True

    verify = client.post("/api/verify", json={"original": text, "candidate": repaired, "genre": "professional-report"})
    assert verify.status_code == 200
    assert verify.json()["passes"] is True

    exported = client.post("/api/export", json={"text": repaired, "html": f"<p>{repaired}</p>", "filename": "report-clean", "format": "docx"})
    assert exported.status_code == 200
    assert exported.content[:2] == b"PK"


class FakeHandler(BaseHTTPRequestHandler):
    candidate = ""
    def do_POST(self):
        length = int(self.headers.get("content-length", "0"))
        body = self.rfile.read(length)
        parsed = json.loads(body.decode())
        assert parsed["model"] == "test-model"
        assert self.headers.get("Authorization") == "Bearer test-secret"
        payload = {"choices": [{"message": {"content": self.candidate}}]}
        raw = json.dumps(payload).encode()
        self.send_response(200); self.send_header("Content-Type", "application/json"); self.send_header("Content-Length", str(len(raw))); self.end_headers(); self.wfile.write(raw)
    def log_message(self, *_args):
        return


def test_byok_rewrite_hits_provider_without_persisting_key():
    baseline = studio.repair_text(SLOP, "professional-report")
    FakeHandler.candidate = baseline["finalText"]
    server = HTTPServer(("127.0.0.1", 0), FakeHandler)
    thread = threading.Thread(target=server.serve_forever, daemon=True); thread.start()
    try:
        response = client.post(
            "/api/rewrite",
            headers={"X-WriteRoute-Key": "test-secret"},
            json={
                "text": SLOP,
                "genre": "professional-report",
                "provider": "openai-compatible",
                "model": "test-model",
                "base_url": f"http://127.0.0.1:{server.server_port}",
                "candidates": 1,
                "temperature": 0.2,
            },
        )
        assert response.status_code == 200, response.text
        result = response.json()
        assert result["changed"] is True
        assert result["finalText"] == baseline["finalText"]
        assert any(a["accepted"] for a in result["attempts"])
    finally:
        server.shutdown(); thread.join(timeout=2)

    # The product has no persistence route and the key never appears in source outputs.
    assert "test-secret" not in json.dumps(client.get("/api/health").json())


def test_rejects_unknown_binary_upload():
    r = client.post("/api/upload", data={"genre": "auto"}, files={"file": ("x.bin", b"abc", "application/octet-stream")})
    assert r.status_code == 415

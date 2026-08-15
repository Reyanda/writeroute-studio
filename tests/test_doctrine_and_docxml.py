from __future__ import annotations
import io
import pytest
from fastapi.testclient import TestClient
from docx import Document
from lxml import etree

from app import app
from writeroute.docxml_engine import docx_to_rich_html, DocxPackage
from writeroute.mendeley_citations import create_mendeley_sdt, create_mendeley_bibliography_marker
from writeroute.auctor_doctrine import (
    FactLedger,
    ThreeChannelEnforcer,
    RevisionAuthority,
    document_artifact_router,
)

client = TestClient(app)


def make_test_docx() -> bytes:
    """Creates a sample docx with headings, bold text, and a table."""
    doc = Document()
    doc.add_heading("Neonatal Clinical Outcomes", level=1)
    p = doc.add_paragraph("The prospective cohort enrolled ")
    r = p.add_run("1,200 infants")
    r.bold = True
    p.add_run(" across 14 health centers.")
    
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "Group"
    t.cell(0, 1).text = "Odds Ratio"
    t.cell(1, 0).text = "Intervention"
    t.cell(1, 1).text = "0.58 (95% CI: 0.44 to 0.76)"
    
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_docxml_engine_rich_html():
    docx_bytes = make_test_docx()
    res = docx_to_rich_html(docx_bytes)
    assert "<h1>" in res["html"]
    assert "Neonatal Clinical Outcomes" in res["html"]
    assert "<strong>" in res["html"]
    assert "1,200 infants" in res["html"]
    assert "<table" in res["html"]
    assert "0.58" in res["html"]


def test_mendeley_citations_sdt():
    sdt = create_mendeley_sdt(
        citation_text="(Smith et al., 2024)",
        items=[{
            "title": "Neonatal Survival",
            "journal": "The Lancet",
            "year": 2024,
            "doi": "10.1016/S0140-6736(24)00123-4",
        }],
    )
    xml_str = etree.tostring(sdt, encoding="unicode")
    assert "w:sdt" in xml_str
    assert "MENDELEY_CITATION_v3_" in xml_str
    assert "(Smith et al., 2024)" in xml_str

    bib_marker = create_mendeley_bibliography_marker()
    bib_str = etree.tostring(bib_marker, encoding="unicode")
    assert "MENDELEY_BIBLIOGRAPHY" in bib_str


def test_fact_ledger_invariants():
    orig_text = "The adjusted odds ratio was 0.58 (95% CI: 0.44 to 0.76; p < 0.001) in 1,200 infants. No adverse events occurred."
    ledger = FactLedger.extract_from_text(orig_text)
    
    assert "0.58" in ledger.numbers
    assert "1,200" in ledger.numbers
    assert len(ledger.confidence_intervals) == 1
    assert len(ledger.p_values) == 1
    assert "no" in [n.lower() for n in ledger.negations]

    # Valid rewrite that preserves facts
    good_rewrite = "In a cohort of 1,200 infants, the adjusted odds ratio was 0.58 (95% CI: 0.44 to 0.76; p < 0.001). No adverse events occurred."
    violations_good = ledger.verify_invariants(good_rewrite)
    assert len(violations_good) == 0

    # Bad rewrite that alters number (0.58 -> 0.85) and removes negation
    bad_rewrite = "In 1,200 infants, the adjusted odds ratio was 0.85 (95% CI: 0.44 to 0.76; p < 0.001). Adverse events occurred."
    violations_bad = ledger.verify_invariants(bad_rewrite)
    assert any("0.58" in v for v in violations_bad)
    assert any("Negation removed" in v for v in violations_bad)


def test_three_channel_enforcer():
    leaky_text = "The study enrolled 1,200 infants. [STATS-BRAIN: SB-CORE-002] severity: critical. Here is the revised text: all patients recovered."
    leaks = ThreeChannelEnforcer.check_channel_leakage(leaky_text)
    assert len(leaks) > 0

    sanitized = ThreeChannelEnforcer.sanitize_substantive_channel(leaky_text)
    assert "STATS-BRAIN" not in sanitized
    assert "severity: critical" not in sanitized
    assert "Here is the revised text" not in sanitized
    assert "1,200 infants" in sanitized


def test_revision_authority_and_router():
    orig = "The rate decreased by 4.2%."
    
    # Mechanical authority fails if words change
    mech_fail = RevisionAuthority.validate_action(
        RevisionAuthority.MECHANICAL,
        orig,
        "The incidence decreased by 4.2%.",
    )
    assert mech_fail["valid"] is False

    # Developmental authority fails if text modified
    dev_fail = RevisionAuthority.validate_action(
        RevisionAuthority.DEVELOPMENTAL,
        orig,
        "The rate reduced by 4.2%.",
    )
    assert dev_fail["valid"] is False

    # Router checks
    docx_route = document_artifact_router("docx")
    assert docx_route["lead_skill"] == "docxml-orchestration"
    assert docx_route["citation_engine"] == "mendeley-citations"


def test_api_doctrine_endpoints():
    docx_bytes = make_test_docx()
    
    # /api/docx/import-rich
    res_import = client.post(
        "/api/docx/import-rich",
        files={"file": ("sample.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert res_import.status_code == 200
    data = res_import.json()
    assert "<h1>" in data["html"]
    assert "1,200 infants" in data["html"]

    # /api/auctor/doctrine-audit
    res_audit = client.post(
        "/api/auctor/doctrine-audit",
        json={
            "original_text": "The adjusted odds ratio was 0.58 (95% CI: 0.44 to 0.76) in 1,200 infants. No adverse events observed.",
            "candidate_text": "In 1,200 infants, the adjusted odds ratio was 0.58 (95% CI: 0.44 to 0.76). No adverse events observed.",
            "authority": "substantive",
        },
    )
    assert res_audit.status_code == 200
    audit_data = res_audit.json()
    assert audit_data["fact_ledger"]["numbers_count"] > 0
    assert audit_data["three_channel_status"]["is_clean"] is True
    assert audit_data["authority_validation"]["valid"] is True

    # /api/router/lead-skill
    res_router = client.get("/api/router/lead-skill?artifact=docx")
    assert res_router.status_code == 200
    assert res_router.json()["lead_skill"] == "docxml-orchestration"

from __future__ import annotations

import json
from pathlib import Path
from typing import Any

import yaml
from docx import Document


TEXT_SUFFIXES = {".txt", ".md", ".markdown", ".rst"}
STRUCTURED_SUFFIXES = {".json", ".yaml", ".yml"}


def load_review_input(path: str | Path) -> dict[str, Any]:
    source = Path(path)
    if not source.exists():
        raise FileNotFoundError(source)
    suffix = source.suffix.lower()
    if suffix in STRUCTURED_SUFFIXES:
        if suffix == ".json":
            value = json.loads(source.read_text(encoding="utf-8"))
        else:
            value = yaml.safe_load(source.read_text(encoding="utf-8"))
        if not isinstance(value, dict):
            raise ValueError("Structured review input must contain an object or mapping")
        value.setdefault("source_name", source.name)
        return value
    if suffix in TEXT_SUFFIXES:
        return {"manifest": {}, "manuscript_text": source.read_text(encoding="utf-8"), "source_name": source.name}
    if suffix == ".docx":
        text, inventory = extract_docx_text(source)
        return {
            "manifest": {"document": {"format": "docx", "paragraph_inventory": inventory}},
            "manuscript_text": text,
            "source_name": source.name,
        }
    raise ValueError(f"Unsupported input type: {suffix}")


def extract_docx_text(path: str | Path) -> tuple[str, list[dict[str, Any]]]:
    document = Document(str(path))
    lines: list[str] = []
    inventory: list[dict[str, Any]] = []
    current_section = "front_matter"
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        style_name = paragraph.style.name if paragraph.style is not None else ""
        if not text:
            continue
        if style_name.lower().startswith("heading"):
            current_section = text
        lines.append(text)
        inventory.append(
            {
                "paragraph_index": index,
                "style": style_name,
                "section": current_section,
                "text": text,
            }
        )
    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            line = " | ".join(cells)
            if line.strip(" |"):
                lines.append(line)
                inventory.append(
                    {
                        "table_index": table_index,
                        "row_index": row_index,
                        "style": "table",
                        "section": "tables",
                        "text": line,
                    }
                )
    return "\n\n".join(lines), inventory
